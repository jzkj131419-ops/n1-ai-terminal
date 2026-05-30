#!/usr/bin/env python3
"""动态银行流水多维度分析。
读取已生成的审计附表 Excel，输出：
1. 在原 Excel 上增加分析 sheet
2. 生成 Word 审计报告
"""
from __future__ import annotations
import os, re
from collections import defaultdict, Counter
from datetime import date
from decimal import Decimal
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_BASE = os.environ.get('N1_OUTPUT_DIR', '/Users/lin/Desktop')
XLSX = Path(os.environ.get('N1_AUDIT_XLSX', f"{OUT_BASE}/审计附表.xlsx"))
DOCX = Path(os.environ.get('N1_AUDIT_DOCX', f"{OUT_BASE}/审计报告.docx"))
MD   = Path(os.environ.get('N1_AUDIT_MD', f"{OUT_BASE}/审计报告.md"))

CUSTOMER = os.environ.get('N1_CUSTOMER_NAME', '未命名客户')
ACCOUNT  = os.environ.get('N1_ACCOUNT_NO', '')
BANK     = os.environ.get('N1_BANK_NAME', '未指定银行')
REPORT_MODE = os.environ.get('N1_REPORT_MODE', 'single_account')
EXTENDED_MODE = os.environ.get('N1_EXTENDED_REPORT', '0') == '1'
SUBJECT_SCOPE_RAW = os.environ.get('N1_SUBJECT_SCOPE', '')
ACCOUNT_SCOPE = os.environ.get('N1_ACCOUNT_SCOPE', '')
BANK_SCOPE = os.environ.get('N1_BANK_SCOPE', '')

# ---------- 1. 读数据 ----------
wb = load_workbook(XLSX)
ws = wb["1交易记录汇总"]
if not CUSTOMER:
    CUSTOMER = str(ws.cell(2, 3).value or '未命名客户')
if not ACCOUNT:
    ACCOUNT = str(ws.cell(2, 5).value or '')
if CUSTOMER in {'审计案件', '当前任务', '未命名客户', ''}:
    CUSTOMER = str(ws.cell(2, 3).value or CUSTOMER or '未命名客户')


def parse_scope(raw: str) -> list[str]:
    return [item.strip() for item in str(raw or "").split("|") if item and item.strip()]


def normalize_subject_name(value: str) -> str:
    text = str(value or "").strip()
    if text in {'审计案件', '当前任务', '未命名客户', '单案件处理', '未填写客户名称', ''}:
        return ''
    return text


AUTO_INFERENCE_TAGS = ("[账号反查]", "[摘要推断]", "[未识别对手]", "[推断:")
AUTO_INFERENCE_FILL = PatternFill("solid", fgColor="FFD580")
NO_FILL = PatternFill(fill_type=None)


def apply_auto_inference_highlights() -> None:
    for row_idx in range(2, ws.max_row + 1):
        cp_acct_cell = ws.cell(row_idx, 11)
        cp_name_cell = ws.cell(row_idx, 12)
        remark_cell = ws.cell(row_idx, 16)
        cp_name = str(cp_name_cell.value or "").strip()
        remark = str(remark_cell.value or "").strip()
        is_auto = any(tag in cp_name for tag in AUTO_INFERENCE_TAGS) or any(tag in remark for tag in AUTO_INFERENCE_TAGS)
        cp_acct_cell.fill = AUTO_INFERENCE_FILL if is_auto else NO_FILL
        cp_name_cell.fill = AUTO_INFERENCE_FILL if is_auto else NO_FILL


apply_auto_inference_highlights()


def infer_subjects_from_workbook() -> list[str]:
    subjects = []
    seen = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        subject = normalize_subject_name(row[2] if len(row) > 2 else "")
        if subject and subject not in seen:
            seen.add(subject)
            subjects.append(subject)
    return subjects

def infer_bank_name_from_workbook():
    names = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        summary = (row[13] or "") if len(row) > 13 else ""
        text_summary = (row[14] or "") if len(row) > 14 else ""
        remark = (row[15] or "") if len(row) > 15 else ""
        joined = " ".join(str(x) for x in [summary, text_summary, remark] if x)
        m = re.search(r"([一-龥A-Za-z]{2,20}银行)", joined)
        if m:
            names.append(m.group(1))
    if names:
        return Counter(names).most_common(1)[0][0]
    stem = XLSX.stem
    m = re.search(r"([一-龥A-Za-z]{2,20}银行)", stem)
    return m.group(1) if m else ""

if not BANK or BANK == '未指定银行':
    inferred_bank = infer_bank_name_from_workbook()
    if inferred_bank:
        BANK = inferred_bank

SUBJECT_SCOPE = parse_scope(SUBJECT_SCOPE_RAW) or infer_subjects_from_workbook()
ACCOUNT_SCOPE_LIST = parse_scope(ACCOUNT_SCOPE)
BANK_SCOPE_LIST = parse_scope(BANK_SCOPE)
if not SUBJECT_SCOPE and normalize_subject_name(CUSTOMER):
    SUBJECT_SCOPE = [CUSTOMER]
if not ACCOUNT_SCOPE_LIST and ACCOUNT:
    ACCOUNT_SCOPE_LIST = [ACCOUNT]
if not BANK_SCOPE_LIST and BANK and BANK != '未指定银行':
    BANK_SCOPE_LIST = [BANK]

if REPORT_MODE != "multi_account" and len(ACCOUNT_SCOPE_LIST) > 1:
    REPORT_MODE = "multi_account"

if REPORT_MODE == "multi_account" and not normalize_subject_name(CUSTOMER):
    if SUBJECT_SCOPE:
        CUSTOMER = SUBJECT_SCOPE[0]
    else:
        CUSTOMER = "待确认主体"

ASSOCIATED_SUBJECTS = [subject for subject in SUBJECT_SCOPE if subject != CUSTOMER]
ASSOCIATED_SUBJECTS_TEXT = "；".join(ASSOCIATED_SUBJECTS) if ASSOCIATED_SUBJECTS else "无"

def write_empty_outputs(reason: str) -> None:
    def write_sheet(name, headers, rows, header_fill="FFD966"):
        if name in wb.sheetnames:
            del wb[name]
        s = wb.create_sheet(name)
        bold = Font(bold=True)
        fill = PatternFill("solid", fgColor=header_fill)
        for ci, h in enumerate(headers, 1):
            c = s.cell(1, ci, h)
            c.font = bold
            c.fill = fill
            c.alignment = Alignment(horizontal="center")
        for ri, row in enumerate(rows, 2):
            for ci, v in enumerate(row, 1):
                s.cell(ri, ci, v)
        if not rows:
            s.cell(2, 1, reason)
            if len(headers) > 1:
                s.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
            s.cell(2, 1).font = Font(italic=True, color="808080")
            s.cell(2, 1).alignment = Alignment(horizontal="left")
        return s

    if EXTENDED_MODE:
        write_sheet("4整体概览", ["项目", "数值"], [
            ["客户名称", CUSTOMER],
            ["关联主体清单", ASSOCIATED_SUBJECTS_TEXT],
            ["开户行", BANK],
            ["账号", ACCOUNT],
            ["状态", "未识别到有效交易"],
            ["说明", reason],
        ])
        for sheet_name in (
            "5年度收支", "6月度收支", "7TOP公司流入", "8TOP公司流出", "9TOP个人流入",
            "10TOP个人流出", "11净流入TOP", "12净流出TOP", "13大额交易(≥50万)",
            "14整10万元整数交易", "15现金存取", "16高频对手(≥50笔)", "17冲正交易"
        ):
            write_sheet(sheet_name, ["说明"], [])
    else:
        for _n in list(wb.sheetnames):
            if _n[:1] in "456789" or _n[:2] in ["10","11","12","13","14","15","16","17"]:
                del wb[_n]
    wb.save(XLSX)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "宋体"
    style.font.size = Pt(11)
    doc.add_heading("银行流水审计报告", 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"被审计单位：{CUSTOMER}").bold = True
    doc.add_heading("一、处理结果", 1)
    doc.add_paragraph(
        f"本次材料未识别到可用于生成审计报告的有效交易记录。\n"
        f"开户行：{BANK}\n"
        f"账号：{ACCOUNT}\n"
        f"原因：{reason}\n"
        f"建议：请优先提供标准银行 CSV、完整银行流水 PDF，或确认压缩包已正确解密后重新处理。"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"报告日期：{date.today().isoformat()}").bold = True
    doc.save(DOCX)
    print(f"已生成 Word 报告: {DOCX}")

    md_lines = [
        "# 银行流水审计报告\n",
        f"**被审计单位：{CUSTOMER}**\n",
        f"**开户行：{BANK}　账号：{ACCOUNT}**\n",
        f"**报告日期：{date.today().isoformat()}**\n",
        "\n## 一、处理结果\n",
        f"- 状态：未识别到有效交易\n",
        f"- 原因：{reason}\n",
        "- 建议：请优先提供标准银行 CSV、完整银行流水 PDF，或确认压缩包已正确解密后重新处理。\n",
    ]
    MD.write_text("\n".join(md_lines), encoding="utf-8")
    print(f"已生成 Markdown 报告: {MD}")
    raise SystemExit(0)

# 列：1日期 2时间 3客户 4卡号 5账号 6方向 7金额 8收入 9支出 10余额 11对手账号 12对手户名 13空 14摘要 15摘要全文 16附言
class Tx:
    __slots__ = ("date","time","direction","amount","income","expense","balance",
                 "cp_acct","cp_name","summary","text_summary","remark","is_reversal")
    def __init__(self, row):
        self.date = row[0] or ""
        self.time = row[1] or ""
        self.direction = row[5] or ""
        self.amount = Decimal(str(row[6])) if row[6] else Decimal("0")
        self.income = Decimal(str(row[7])) if row[7] else Decimal("0")
        self.expense = Decimal(str(row[8])) if row[8] else Decimal("0")
        self.balance = Decimal(str(row[9])) if row[9] else Decimal("0")
        self.cp_acct = (row[10] or "").strip()
        self.cp_name = (row[11] or "").strip()
        self.summary = (row[13] or "").strip()
        self.text_summary = (row[14] or "").strip()
        self.remark = (row[15] or "").strip()
        self.is_reversal = "[冲正]" in self.remark

txs = [Tx(r) for r in ws.iter_rows(min_row=2, values_only=True) if r[0]]
print(f"加载 {len(txs)} 条交易")

if not txs:
    write_empty_outputs("审计附表中未识别到有效交易日期，通常是输入材料为空、CSV 字段不兼容或压缩包未成功解密。")

# 排除冲正及与冲正等额的原扣款（用于对账）
real_txs = [t for t in txs if not t.is_reversal]

# ---------- 2. 整体概览 ----------
total_income = sum(t.income for t in txs)
total_expense = sum(t.expense for t in txs)
income_excl_reversal = sum(t.income for t in txs if not t.is_reversal)
expense_excl_reversal = sum(t.expense for t in txs if not t.is_reversal)
net = total_income - total_expense

dates = sorted({t.date for t in txs})
if not dates:
    write_empty_outputs("审计附表中未识别到有效交易日期，通常是输入材料为空、CSV 字段不兼容或压缩包未成功解密。")
first_date = dates[0]
last_date  = dates[-1]
opening_bal = txs[0].balance - txs[0].income + txs[0].expense
closing_bal = txs[-1].balance

# ---------- 3. 公司 vs 个人分类 ----------
COMPANY_KW = ["公司","中心","银行","股份","有限","集团","厂","店","所","局","部","处","会",
              "学校","医院","保险","基金","证券","税务","事务","工商","住房","物业","商行","支行","分行",
              "合作社","酒店","商城","商贸","贸易","实业","咨询","服务","管理","科技","传媒","文化",
              "投资","控股","建设","工程","装饰","设计","教育","婚庆","摄影","影视","传播","培训"]

def classify_counterpart(name: str, acct: str) -> str:
    if not name and not acct:
        return "未知"
    if not name:
        return "未知"  # 对手名为空但有账号 → 可能是个人/系统
    if any(kw in name for kw in COMPANY_KW):
        return "公司/机构"
    # 移除括号里的内容再判断长度
    pure = re.sub(r"[（(].*?[)）]", "", name).strip()
    if 2 <= len(pure) <= 4 and re.fullmatch(r"[一-龥·]+", pure):
        return "个人"
    return "其他"

# ---------- 4. 按对手聚合 ----------
agg = defaultdict(lambda: {"name":"", "acct":"", "kind":"", "n":0,
                           "income":Decimal("0"), "expense":Decimal("0"),
                           "first":None, "last":None, "summaries":Counter()})
for t in txs:
    if t.is_reversal:
        continue
    key = (t.cp_name, t.cp_acct)
    a = agg[key]
    a["name"] = t.cp_name
    a["acct"] = t.cp_acct
    a["kind"] = classify_counterpart(t.cp_name, t.cp_acct)
    a["n"] += 1
    a["income"] += t.income
    a["expense"] += t.expense
    if a["first"] is None or t.date < a["first"]:
        a["first"] = t.date
    if a["last"] is None or t.date > a["last"]:
        a["last"] = t.date
    if t.summary:
        a["summaries"][t.summary] += 1

agg_list = [v for v in agg.values() if v["name"] or v["acct"]]
for a in agg_list:
    a["net"] = a["income"] - a["expense"]
    a["abs_total"] = a["income"] + a["expense"]

print(f"对手聚合: {len(agg_list)} 个")

# ---------- 5. 年度/月度 ----------
def year_of(d):
    s = str(d)
    m = re.match(r"(\d{4})", s)
    return m.group(1) if m else "?"

def ym_of(d):
    s = str(d)
    m = re.match(r"(\d{4})-(\d{1,2})", s)
    return f"{m.group(1)}-{int(m.group(2)):02d}" if m else "?"

yearly = defaultdict(lambda: {"n":0, "income":Decimal("0"), "expense":Decimal("0")})
monthly = defaultdict(lambda: {"n":0, "income":Decimal("0"), "expense":Decimal("0")})
for t in txs:
    if t.is_reversal: continue
    y = year_of(t.date); m = ym_of(t.date)
    yearly[y]["n"] += 1; yearly[y]["income"] += t.income; yearly[y]["expense"] += t.expense
    monthly[m]["n"] += 1; monthly[m]["income"] += t.income; monthly[m]["expense"] += t.expense

# ---------- 6. 大额交易 ----------
LARGE_THRESHOLD = Decimal("500000")
HUGE_THRESHOLD  = Decimal("1000000")
large_txs = sorted([t for t in txs if (t.income >= LARGE_THRESHOLD or t.expense >= LARGE_THRESHOLD) and not t.is_reversal],
                   key=lambda t: max(t.income, t.expense), reverse=True)

# ---------- 7. 整数额交易（万元整、十万整） ----------
def is_round(v: Decimal, base: Decimal) -> bool:
    if v == 0: return False
    return v % base == 0

round_10w_txs = [t for t in txs if not t.is_reversal and (
    (t.income > 0 and is_round(t.income, Decimal("100000"))) or
    (t.expense > 0 and is_round(t.expense, Decimal("100000")))
)]

# ---------- 8. 现金存取 ----------
CASH_KW = ["现金", "取款", "存款", "ATM", "柜台"]
cash_txs = [t for t in txs if not t.is_reversal and any(k in (t.summary + t.text_summary) for k in CASH_KW)]

# ---------- 9. 高频对手（笔数 ≥ 50） ----------
high_freq = [a for a in agg_list if a["n"] >= 50]
high_freq.sort(key=lambda a: a["n"], reverse=True)

# ---------- 10. TOP 排序 ----------
top_company_in  = sorted([a for a in agg_list if a["kind"]=="公司/机构"], key=lambda a: a["income"], reverse=True)[:20]
top_company_out = sorted([a for a in agg_list if a["kind"]=="公司/机构"], key=lambda a: a["expense"], reverse=True)[:20]
top_person_in   = sorted([a for a in agg_list if a["kind"]=="个人"],     key=lambda a: a["income"], reverse=True)[:20]
top_person_out  = sorted([a for a in agg_list if a["kind"]=="个人"],     key=lambda a: a["expense"], reverse=True)[:20]
top_net_out     = sorted(agg_list, key=lambda a: a["net"])[:20]   # 净流出（负数最大）
top_net_in      = sorted(agg_list, key=lambda a: a["net"], reverse=True)[:20]

# 集中度
total_in = sum(a["income"] for a in agg_list)
total_out = sum(a["expense"] for a in agg_list)
top10_in_share = sum(a["income"] for a in sorted(agg_list, key=lambda a: a["income"], reverse=True)[:10]) / total_in * 100 if total_in else Decimal("0")
top10_out_share = sum(a["expense"] for a in sorted(agg_list, key=lambda a: a["expense"], reverse=True)[:10]) / total_out * 100 if total_out else Decimal("0")

# ---------- 11. 写 Excel sheets ----------
def write_sheet(name, headers, rows, header_fill="FFD966"):
    if name in wb.sheetnames:
        del wb[name]
    s = wb.create_sheet(name)
    bold = Font(bold=True)
    fill = PatternFill("solid", fgColor=header_fill)
    for ci, h in enumerate(headers, 1):
        c = s.cell(1, ci, h)
        c.font = bold
        c.fill = fill
        c.alignment = Alignment(horizontal="center")
    for ri, row in enumerate(rows, 2):
        for ci, v in enumerate(row, 1):
            if isinstance(v, Decimal):
                v = float(v)
            s.cell(ri, ci, v)
    if not rows:
        note = "本案未识别到相关交易。"
        s.cell(2, 1, note)
        if len(headers) > 1:
            s.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
        s.cell(2, 1).font = Font(italic=True, color="808080")
        s.cell(2, 1).alignment = Alignment(horizontal="left")
    # 列宽
    for ci in range(1, len(headers)+1):
        s.column_dimensions[chr(64+ci) if ci<=26 else 'A'+chr(64+ci-26)].width = 15
    return s

# 4 整体概览
if EXTENDED_MODE:
    write_sheet("4整体概览", ["项目","数值"], [
        ["客户名称", CUSTOMER],
        ["关联主体清单", ASSOCIATED_SUBJECTS_TEXT],
        ["报告模式", "多账户汇总" if REPORT_MODE == "multi_account" else "单账户审计"],
        ["开户行", "；".join(BANK_SCOPE_LIST) if REPORT_MODE == "multi_account" and BANK_SCOPE_LIST else BANK],
        ["账号", f"{len(ACCOUNT_SCOPE_LIST)}个账户，详见附表2" if REPORT_MODE == "multi_account" and ACCOUNT_SCOPE_LIST else ACCOUNT],
        ["流水起止", f"{first_date}  至  {last_date}"],
        ["交易笔数(含冲正)", len(txs)],
        ["交易笔数(剔除冲正)", len(txs) - sum(1 for t in txs if t.is_reversal)],
        ["冲正笔数", sum(1 for t in txs if t.is_reversal)],
        ["收入合计(原始解析)", float(total_income)],
        ["支出合计(原始解析)", float(total_expense)],
        ["收入合计(剔除冲正)", float(income_excl_reversal)],
        ["支出合计(剔除冲正)", float(expense_excl_reversal)],
        ["净额", float(net)],
        ["期初余额", float(opening_bal)],
        ["期末余额", float(closing_bal)],
        ["对手总数", len(agg_list)],
        ["TOP10流入对手集中度", f"{float(top10_in_share):.2f}%"],
        ["TOP10流出对手集中度", f"{float(top10_out_share):.2f}%"],
    ])
    
    # 5 年度收支汇总
    write_sheet("5年度收支", ["年份","笔数","收入","支出","净额"],
        [[y, v["n"], float(v["income"]), float(v["expense"]), float(v["income"]-v["expense"])]
         for y, v in sorted(yearly.items())])
    
    # 6 月度收支
    write_sheet("6月度收支", ["月份","笔数","收入","支出","净额"],
        [[m, v["n"], float(v["income"]), float(v["expense"]), float(v["income"]-v["expense"])]
         for m, v in sorted(monthly.items())])
    
    # 7 公司对手TOP 流入
    def fmt_agg(a):
        top_summary = a["summaries"].most_common(1)[0][0] if a["summaries"] else ""
        return [a["name"], a["acct"], a["kind"], a["n"],
                float(a["income"]), float(a["expense"]), float(a["net"]),
                str(a["first"]), str(a["last"]), top_summary]
    
    cols = ["对方户名","对方账号","对手类型","笔数","累计流入","累计流出","净额","首次","末次","摘要码"]
    write_sheet("7TOP公司流入", cols, [fmt_agg(a) for a in top_company_in])
    write_sheet("8TOP公司流出", cols, [fmt_agg(a) for a in top_company_out])
    write_sheet("9TOP个人流入", cols, [fmt_agg(a) for a in top_person_in])
    write_sheet("10TOP个人流出", cols, [fmt_agg(a) for a in top_person_out])
    write_sheet("11净流入TOP", cols, [fmt_agg(a) for a in top_net_in])
    write_sheet("12净流出TOP", cols, [fmt_agg(a) for a in top_net_out])
    
    # 13 大额交易
    write_sheet("13大额交易(≥50万)",
        ["日期","时间","方向","金额","余额","对方户名","对方账号","对手类型","摘要","摘要全文"],
        [[str(t.date), str(t.time), t.direction,
          float(t.income or t.expense),
          float(t.balance), t.cp_name, t.cp_acct,
          classify_counterpart(t.cp_name, t.cp_acct),
          t.summary, t.text_summary]
         for t in large_txs])
    
    # 14 整十万元交易
    write_sheet("14整10万元整数交易",
        ["日期","时间","方向","金额","对方户名","对方账号","摘要","摘要全文"],
        [[str(t.date), str(t.time), t.direction,
          float(t.income or t.expense),
          t.cp_name, t.cp_acct, t.summary, t.text_summary]
         for t in round_10w_txs])
    
    # 15 现金类
    write_sheet("15现金存取",
        ["日期","时间","方向","金额","余额","摘要","摘要全文"],
        [[str(t.date), str(t.time), t.direction,
          float(t.income or t.expense),
          float(t.balance), t.summary, t.text_summary]
         for t in cash_txs])
    
    # 16 高频对手
    write_sheet("16高频对手(≥50笔)",
        cols + ["日均笔数(活跃区间)"],
        [fmt_agg(a) + [round(a["n"] / max(1, (date.fromisoformat(str(a["last"])) - date.fromisoformat(str(a["first"]))).days+1), 3)
                        if str(a["first"]) and str(a["last"]) and "-" in str(a["first"]) else 0]
         for a in high_freq])
    
    # 17 冲正交易
    reversal_txs = [t for t in txs if t.is_reversal]
    write_sheet("17冲正交易",
        ["日期","时间","方向","金额","对方户名","摘要","摘要全文","备注"],
        [[str(t.date), str(t.time), t.direction,
          float(t.income or t.expense),
          t.cp_name, t.summary, t.text_summary, t.remark]
         for t in reversal_txs])
    
else:
    for _n in list(wb.sheetnames):
        if _n[:1] in "456789" or _n[:2] in ["10","11","12","13","14","15","16","17"]:
            del wb[_n]
wb.save(XLSX)
if "Sheet1" in wb.sheetnames:
    del wb["Sheet1"]
    wb.save(XLSX)
print(f"已写入 Excel 分析 sheet 到: {XLSX}")

# ---------- 12. 写 Word 报告 ----------
def yuan(v):
    """金额格式化：保留2位，加千分位。"""
    if isinstance(v, Decimal): v = float(v)
    return f"{v:,.2f}"

doc = Document()
# 字体
style = doc.styles['Normal']
style.font.name = "宋体"
style.font.size = Pt(11)

doc.add_heading("银行流水审计报告", 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"被审计单位：{CUSTOMER}").bold = True

doc.add_heading("一、基本情况", 1)
if REPORT_MODE == "multi_account":
    doc.add_paragraph(
        f"本次审计以 {CUSTOMER} 为主口径，对提交材料中识别出的多个银行账户进行汇总分析。\n"
        f"主体范围：{'；'.join(SUBJECT_SCOPE) if SUBJECT_SCOPE else CUSTOMER}\n"
        f"关联主体清单：{ASSOCIATED_SUBJECTS_TEXT}\n"
        f"账户范围：共 {len(ACCOUNT_SCOPE_LIST) if ACCOUNT_SCOPE_LIST else 0} 个账户，详见附表2《账号汇总清单》\n"
        f"涉及银行：{'；'.join(BANK_SCOPE_LIST) if BANK_SCOPE_LIST else (BANK or '待补充确认')}\n"
        f"流水期间：{first_date} 至 {last_date}\n"
        f"汇总口径期初余额：¥{yuan(opening_bal)}\n"
        f"汇总口径期末余额：¥{yuan(closing_bal)}"
    )
else:
    doc.add_paragraph(
        f"本次审计的银行账户基本情况如下：\n"
        f"账户名称：{CUSTOMER}\n"
        f"关联主体清单：{ASSOCIATED_SUBJECTS_TEXT}\n"
        f"开户行：{BANK}\n"
        f"账号：{ACCOUNT}\n"
        f"流水期间：{first_date} 至 {last_date}\n"
        f"账户期初余额：¥{yuan(opening_bal)}\n"
        f"账户期末余额：¥{yuan(closing_bal)}"
    )

doc.add_heading("二、整体收支概况", 1)
overview_prefix = "汇总流水" if REPORT_MODE == "multi_account" else "原始流水"
doc.add_paragraph(
    f"{overview_prefix}共计 {len(txs):,} 笔交易，其中含 {sum(1 for t in txs if t.is_reversal)} 笔银行系统冲正交易（系银行自行修正、不构成实际资金往来）。"
    f"剔除冲正后实际交易 {len(txs) - sum(1 for t in txs if t.is_reversal):,} 笔。\n\n"
    f"当前口径下，剔除冲正后的汇总结果为：\n"
    f"  · 收入合计：¥{yuan(income_excl_reversal)}\n"
    f"  · 支出合计：¥{yuan(expense_excl_reversal)}\n"
    f"  · 净额    ：¥{yuan(net)}"
)

if EXTENDED_MODE:
    doc.add_heading("三、年度收支分布", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["年份", "笔数", "收入", "支出", "净额"]):
        hdr[i].text = h
    for y, v in sorted(yearly.items()):
        row = table.add_row().cells
        row[0].text = y
        row[1].text = str(v["n"])
        row[2].text = yuan(v["income"])
        row[3].text = yuan(v["expense"])
        row[4].text = yuan(v["income"] - v["expense"])
    
    doc.add_heading("四、对手集中度分析", 1)
    doc.add_paragraph(
        f"账户共涉及 {len(agg_list):,} 个不重复对手（按户名+账号去重）。\n"
        f"  · TOP10 流入对手占总流入比：{float(top10_in_share):.2f}%\n"
        f"  · TOP10 流出对手占总流出比：{float(top10_out_share):.2f}%\n"
    )
    
    doc.add_heading("五、累计流入TOP10（公司/机构）", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["对方户名", "笔数", "累计流入", "累计流出", "净额"]):
        table.rows[0].cells[i].text = h
    for a in top_company_in[:10]:
        row = table.add_row().cells
        row[0].text = a["name"][:30]
        row[1].text = str(a["n"])
        row[2].text = yuan(a["income"])
        row[3].text = yuan(a["expense"])
        row[4].text = yuan(a["net"])
    
    doc.add_heading("六、累计流出TOP10（公司/机构）", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["对方户名", "笔数", "累计流入", "累计流出", "净额"]):
        table.rows[0].cells[i].text = h
    for a in top_company_out[:10]:
        row = table.add_row().cells
        row[0].text = a["name"][:30]
        row[1].text = str(a["n"])
        row[2].text = yuan(a["income"])
        row[3].text = yuan(a["expense"])
        row[4].text = yuan(a["net"])
    
    doc.add_heading("七、个人对手TOP10（按累计流入排序）", 1)
    doc.add_paragraph("【提示】司法案件审计中应重点关注：①与嫌疑人同名或关联人姓名的对手；②高频或大额的个人对手；③与公司业务无明显关联的个人。")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["对方户名", "笔数", "累计流入", "累计流出", "净额"]):
        table.rows[0].cells[i].text = h
    for a in top_person_in[:10]:
        row = table.add_row().cells
        row[0].text = a["name"]
        row[1].text = str(a["n"])
        row[2].text = yuan(a["income"])
        row[3].text = yuan(a["expense"])
        row[4].text = yuan(a["net"])
    
    doc.add_heading("八、个人对手TOP10（按累计流出排序）", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["对方户名", "笔数", "累计流入", "累计流出", "净额"]):
        table.rows[0].cells[i].text = h
    for a in top_person_out[:10]:
        row = table.add_row().cells
        row[0].text = a["name"]
        row[1].text = str(a["n"])
        row[2].text = yuan(a["income"])
        row[3].text = yuan(a["expense"])
        row[4].text = yuan(a["net"])
    
    doc.add_heading("九、大额交易（单笔≥50万元）", 1)
    doc.add_paragraph(f"剔除冲正后共计 {len(large_txs)} 笔大额交易。详见附表13。前10条摘录：")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["日期", "方向", "金额", "对方户名", "摘要"]):
        table.rows[0].cells[i].text = h
    for t in large_txs[:10]:
        row = table.add_row().cells
        row[0].text = str(t.date)
        row[1].text = t.direction
        row[2].text = yuan(t.income or t.expense)
        row[3].text = t.cp_name[:25]
        row[4].text = (t.summary or t.text_summary)[:30]
    
    doc.add_heading("十、整数额交易（10万元整数倍）", 1)
    doc.add_paragraph(
        f"账户中按10万元整数倍发生的交易共 {len(round_10w_txs)} 笔。司法案件中，此类整数额"
        f"往往与正常业务交易（含税尾数）的资金特征不同，需结合具体业务背景核查。详见附表14。"
    )
    
    doc.add_heading("十一、现金类交易", 1)
    if cash_txs:
        cash_in = sum(t.income for t in cash_txs)
        cash_out = sum(t.expense for t in cash_txs)
        doc.add_paragraph(f"涉及现金存取交易 {len(cash_txs)} 笔，现金流入 ¥{yuan(cash_in)}，现金流出 ¥{yuan(cash_out)}。详见附表15。")
    else:
        doc.add_paragraph("本账户内未识别到明显的现金存取类交易（摘要中含'现金/取款/存款'关键字）。")
    
    doc.add_heading("十二、高频往来对手（≥50笔）", 1)
    doc.add_paragraph(f"与该账户发生 50 笔以上交易的对手共 {len(high_freq)} 个。该类对手通常为长期业务伙伴或固定结算对象，详见附表16。前10位：")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["对方户名", "类型", "笔数", "累计流入", "累计流出"]):
        table.rows[0].cells[i].text = h
    for a in high_freq[:10]:
        row = table.add_row().cells
        row[0].text = a["name"][:30]
        row[1].text = a["kind"]
        row[2].text = str(a["n"])
        row[3].text = yuan(a["income"])
        row[4].text = yuan(a["expense"])
    
    doc.add_heading("十三、冲正交易", 1)
    doc.add_paragraph(
        f"本账户共发生 {sum(1 for t in txs if t.is_reversal)} 笔冲正交易，"
        f"合计金额 ¥{yuan(sum(t.income for t in txs if t.is_reversal))}。"
        "冲正系银行系统对原误操作的反向修正，对客户实际资产无影响，但其对应的原始交易仍存在。详见附表17。"
    )
    
doc.add_heading("十四、审计提示与建议", 1)
advice_intro = (
    "本报告基于多个银行账户的汇总口径生成，为后续案件定性提供资金侧线索。建议先按附表2核定账户范围，再结合以下补充资料进一步定向核查：\n"
    if REPORT_MODE == "multi_account"
    else "本报告基于单一银行账户流水生成，为后续案件定性提供资金侧线索。建议结合以下补充资料进一步定向核查：\n"
)
doc.add_paragraph(
    advice_intro +
    "1. 嫌疑人及关联人姓名清单：用于在对手户名中精确匹配；\n"
    "2. 嫌疑人微信支付/支付宝账单：用于补充小额、移动端资金流；\n"
    "3. 嫌疑人或本案关联公司其他银行账户流水：用于资金链路延伸追踪；\n"
    "4. 业务合同、聊天记录、出入库单据：用于核查大额交易的真实业务背景；\n"
    "5. 嫌疑人个人消费凭证（购车/购房/旅游/奢侈品等）：用于核查个人挪用资金的最终去向。\n\n"
    "依据上述补充资料，应进一步开展：①将嫌疑人姓名与对手户名做穷举匹配；②对个人对手按金额阈值筛查；"
    "③对整数额、深夜、节假日等异常时间窗口的交易做单笔追溯；④识别可能的同名多账户、关联账户串联。"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run(f"报告日期：{date.today().isoformat()}").bold = True

doc.save(DOCX)
print(f"已生成 Word 报告: {DOCX}")

# 同步生成 Markdown 版（便于快速查看）
def md_table(headers, rows):
    out = "| " + " | ".join(headers) + " |\n"
    out += "| " + " | ".join("---" for _ in headers) + " |\n"
    for r in rows:
        out += "| " + " | ".join(str(c) for c in r) + " |\n"
    return out

md_lines = []
md_lines.append(f"# 银行流水审计报告\n")
md_lines.append(f"**被审计单位：{CUSTOMER}**\n")
if REPORT_MODE == "multi_account":
    md_lines.append(f"**报告模式：多账户汇总**\n")
    md_lines.append(f"**涉及银行：{'；'.join(BANK_SCOPE_LIST) if BANK_SCOPE_LIST else (BANK or '待补充确认')}**\n")
    md_lines.append(f"**账户范围：共 {len(ACCOUNT_SCOPE_LIST) if ACCOUNT_SCOPE_LIST else 0} 个账户，详见附表2**\n")
else:
    md_lines.append(f"**开户行：{BANK}　账号：{ACCOUNT}**\n")
md_lines.append(f"**流水期间：{first_date} 至 {last_date}**\n")
md_lines.append(f"**报告日期：{date.today().isoformat()}**\n\n---\n")

md_lines.append("## 一、基本情况\n")
md_lines.append(f"- 关联主体清单：{ASSOCIATED_SUBJECTS_TEXT}")
md_lines.append(f"- 账户期初余额：¥{yuan(opening_bal)}")
md_lines.append(f"- 账户期末余额：¥{yuan(closing_bal)}\n")

md_lines.append("## 二、整体收支概况\n")
md_lines.append(f"- 原始流水：{len(txs):,} 笔（含冲正 {sum(1 for t in txs if t.is_reversal)} 笔）")
md_lines.append(f"- 剔除冲正：{len(txs) - sum(1 for t in txs if t.is_reversal):,} 笔")
md_lines.append(f"- 收入合计（剔除冲正）：¥{yuan(income_excl_reversal)}")
md_lines.append(f"- 支出合计（剔除冲正）：¥{yuan(expense_excl_reversal)}")
md_lines.append(f"- 净额：¥{yuan(net)}（资金净流出）\n")
md_lines.append("上述金额为当前识别口径下的汇总结果，若客户另附银行官方合计页，应再做逐项核对。\n")

if EXTENDED_MODE:
    md_lines.append("## 三、年度收支分布\n")
    md_lines.append(md_table(["年份","笔数","收入","支出","净额"],
        [[y, v["n"], yuan(v["income"]), yuan(v["expense"]), yuan(v["income"]-v["expense"])]
         for y, v in sorted(yearly.items())]))
    
    md_lines.append("\n## 四、对手集中度\n")
    md_lines.append(f"- 对手总数：{len(agg_list):,}")
    md_lines.append(f"- TOP10流入集中度：{float(top10_in_share):.2f}%")
    md_lines.append(f"- TOP10流出集中度：{float(top10_out_share):.2f}%\n")
    
    md_lines.append("## 五、TOP10公司/机构（按累计流入）\n")
    md_lines.append(md_table(["对方户名","笔数","累计流入","累计流出","净额"],
        [[a["name"][:30], a["n"], yuan(a["income"]), yuan(a["expense"]), yuan(a["net"])]
         for a in top_company_in[:10]]))
    
    md_lines.append("\n## 六、TOP10公司/机构（按累计流出）\n")
    md_lines.append(md_table(["对方户名","笔数","累计流入","累计流出","净额"],
        [[a["name"][:30], a["n"], yuan(a["income"]), yuan(a["expense"]), yuan(a["net"])]
         for a in top_company_out[:10]]))
    
    md_lines.append("\n## 七、TOP10个人（按累计流入）\n")
    md_lines.append(md_table(["对方户名","笔数","累计流入","累计流出","净额"],
        [[a["name"], a["n"], yuan(a["income"]), yuan(a["expense"]), yuan(a["net"])]
         for a in top_person_in[:10]]))
    
    md_lines.append("\n## 八、TOP10个人（按累计流出）\n")
    md_lines.append(md_table(["对方户名","笔数","累计流入","累计流出","净额"],
        [[a["name"], a["n"], yuan(a["income"]), yuan(a["expense"]), yuan(a["net"])]
         for a in top_person_out[:10]]))
    
    md_lines.append(f"\n## 九、大额交易（≥50万元）\n")
    md_lines.append(f"剔除冲正后 {len(large_txs)} 笔。详见Excel附表13。\n")
    
    md_lines.append(f"\n## 十、整10万元整数交易\n")
    md_lines.append(f"共 {len(round_10w_txs)} 笔。详见Excel附表14。\n")
    
    md_lines.append(f"\n## 十一、现金类交易\n")
    if cash_txs:
        md_lines.append(f"共 {len(cash_txs)} 笔，流入 ¥{yuan(sum(t.income for t in cash_txs))}，流出 ¥{yuan(sum(t.expense for t in cash_txs))}。详见Excel附表15。\n")
    else:
        md_lines.append("未识别到明显的现金存取类交易。\n")
    
    md_lines.append(f"\n## 十二、高频对手（≥50笔）\n")
    md_lines.append(f"共 {len(high_freq)} 个。详见Excel附表16。\n")
    
    md_lines.append(f"\n## 十三、冲正交易\n")
    md_lines.append(f"共 {sum(1 for t in txs if t.is_reversal)} 笔，金额合计 ¥{yuan(sum(t.income for t in txs if t.is_reversal))}。详见Excel附表17。\n")
    
md_lines.append(f"\n## 十四、审计提示与建议\n")
md_lines.append(
    "本报告基于单一银行账户流水生成。后续定向核查需补充：\n"
    "1. 嫌疑人及关联人姓名清单 → 在对手户名中穷举匹配\n"
    "2. 微信支付 / 支付宝账单 → 补充小额移动端资金流\n"
    "3. 嫌疑人/关联公司其他银行账户流水 → 资金链路延伸追踪\n"
    "4. 业务合同、聊天记录、出入库单据 → 核查大额交易真实业务背景\n"
    "5. 个人消费凭证（购车/购房/旅游/奢侈品） → 追溯个人挪用资金最终去向\n"
)

MD.write_text("\n".join(md_lines), encoding="utf-8")
print(f"已生成 Markdown 报告: {MD}")
